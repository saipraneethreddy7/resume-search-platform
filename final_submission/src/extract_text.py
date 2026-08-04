"""
extract_text.py
----------------
Phase 1 utility: pulls raw text out of resume files (.pdf / .docx) so it can
be handed to the LLM parser. Kept deliberately dependency-light (pdfplumber +
python-docx) so it runs anywhere without LibreOffice/pandoc installed --
important for "design for scalability" since this will eventually run in a
batch/serverless pipeline, not just a notebook kernel.
"""

from pathlib import Path
import sys
import docx
import pdfplumber


def extract_text_from_pdf(path: Path) -> str:
    """Concatenate text across all pages of a PDF resume."""
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                chunks.append(page_text)
    text = "\n".join(chunks)
    if len(text.strip()) < 20:
        print(
            f"WARNING — {path.name} returned no text, this may be a scanned "
            "image PDF. Consider running OCR on this file."
        )
    return text


def extract_text_from_docx(path: Path) -> str:
    """
    Pull text from a .docx, including text sitting inside tables --
    a lot of these resumes format work-experience entries as tables
    (Organization | Duration / Designation | Key Role), so a naive
    paragraph-only read silently drops most of the content.
    """
    document = docx.Document(path)
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            deduped = []
            for c in cells:
                if c and (not deduped or c != deduped[-1]):
                    deduped.append(c)
            if deduped:
                parts.append(" | ".join(deduped))

    return "\n".join(parts)


def extract_text(path: Path) -> str:
    """Dispatch on file extension. Raises for unsupported types."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in (".docx", ".dotx"):
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported resume file type: {suffix} ({path.name})")


def load_all_resumes(folder: Path) -> dict[str, str]:
    """
    Read every supported resume file in `folder` and return
    {filename: raw_text}. Skips files that fail extraction but logs why,
    rather than crashing the whole batch -- at scale, a handful of corrupt
    or oddly-formatted files shouldn't take down the run.
    """
    texts = {}
    for path in sorted(folder.iterdir()):
        if path.suffix.lower() not in (".pdf", ".docx", ".dotx"):
            continue
        try:
            texts[path.name] = extract_text(path)
        except Exception as e:
            print(f"[extract_text] FAILED on {path.name}: {e}")
    return texts


if __name__ == "__main__" and "ipykernel" not in sys.modules:
    # Only run as a CLI tool (`python extract_text.py <folder>`) -- guarded against
    # firing inside Jupyter, where every cell also executes with __name__ == "__main__"
    # and sys.argv holds the kernel's own connection-file args, not a folder path.
    folder = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("../")
    resumes = load_all_resumes(folder)
    for name, text in resumes.items():
        print(f"--- {name} ({len(text)} chars) ---")
